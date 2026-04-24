"""
Phase 4.5: DOCX Report Generator
===================================
Generates a formatted Word document (.docx) containing:

  Page 1:   Color legend + evidence summary table
  Page 2-3: Original clinical note with color-highlighted evidence spans
  Page 4+:  Full clinical preparedness report with branch-colored headers

Color scheme:
  Red    (FF4444) — Branch A: Psychiatric/Clinical Deterioration
  Orange (FF8C00) — Branch B: Substance Use Escalation
  Blue   (4488FF) — Branch C: Social/Environmental Collapse
  Yellow (FFD700) — Nexus Factor: supports 2+ branches

Requires: python-docx (pip install python-docx)
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import List, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

log = logging.getLogger(__name__)


# ============================================================================
# COLOR DEFINITIONS
# ============================================================================

BRANCH_COLORS = {
    "A": {"rgb": RGBColor(0xFF, 0x44, 0x44) if HAS_DOCX else None, "hex": "FF4444", "name": "Red", "label": "Psychiatric / Clinical Deterioration"},
    "B": {"rgb": RGBColor(0xFF, 0x8C, 0x00) if HAS_DOCX else None, "hex": "FF8C00", "name": "Orange", "label": "Substance Use Escalation"},
    "C": {"rgb": RGBColor(0x44, 0x88, 0xFF) if HAS_DOCX else None, "hex": "4488FF", "name": "Blue", "label": "Social / Environmental Collapse"},
    "nexus": {"rgb": RGBColor(0xFF, 0xD7, 0x00) if HAS_DOCX else None, "hex": "FFD700", "name": "Yellow", "label": "Nexus Factor (Multi-Branch)"},
}

HEADER_BLUE = RGBColor(0x1F, 0x4E, 0x79) if HAS_DOCX else None


# ============================================================================
# HELPERS
# ============================================================================


def _add_highlight(run, color_hex: str):
    """Apply a background highlight color to a run using XML manipulation."""
    if not HAS_DOCX:
        return
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    rPr.append(shd)


def _set_cell_shading(cell, color_hex: str):
    """Set background shading on a table cell."""
    if not HAS_DOCX:
        return
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def _add_section_header(doc, text: str, level: int = 2, color: Optional[object] = None):
    """Add a styled section header."""
    heading = doc.add_heading(text, level=level)
    if color and HAS_DOCX:
        for run in heading.runs:
            run.font.color.rgb = color


def _add_para(doc, text: str, bold: bool = False, size: int = 10):
    """Add a paragraph with optional styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    return p


# ============================================================================
# PAGE 1: LEGEND & EVIDENCE SUMMARY
# ============================================================================


def _build_legend_page(doc, evidence_result: dict, silver_labels: dict, report: dict):
    """Build the color legend and evidence summary table."""

    # Title
    title = doc.add_heading("Clinical Preparedness Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("with Counterfactual Evidence Attribution (XAI)")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Metadata
    dx = silver_labels.get("primary_mh_diagnosis", {})
    dx_title = dx.get("title", "Unknown") if isinstance(dx, dict) else "Unknown"
    risk_tier = report.get("overall_risk_tier", "?")

    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").font.size = Pt(9)
    meta.add_run(f"  |  Primary Dx: {dx_title}").font.size = Pt(9)
    meta.add_run(f"  |  Risk Tier: {risk_tier}").font.size = Pt(9)

    doc.add_paragraph()  # spacer

    # --- Color Legend ---
    _add_section_header(doc, "Evidence Color Legend", level=2, color=HEADER_BLUE)

    legend_table = doc.add_table(rows=5, cols=3)
    legend_table.style = 'Table Grid'
    legend_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["Color", "Branch", "Meaning"]
    for i, h in enumerate(headers):
        cell = legend_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    # Data rows
    legend_data = [
        ("A", "Red — Branch A", "Psychiatric / Clinical Deterioration"),
        ("B", "Orange — Branch B", "Substance Use Escalation"),
        ("C", "Blue — Branch C", "Social / Environmental Collapse"),
        ("nexus", "Yellow — Nexus", "Supports 2+ branches (highest priority)"),
    ]
    for row_idx, (key, color_name, meaning) in enumerate(legend_data, 1):
        color_cell = legend_table.rows[row_idx].cells[0]
        _set_cell_shading(color_cell, BRANCH_COLORS[key]["hex"])
        color_cell.text = ""

        legend_table.rows[row_idx].cells[1].text = color_name
        legend_table.rows[row_idx].cells[2].text = meaning

    doc.add_paragraph()  # spacer

    # --- Evidence Summary Table ---
    _add_section_header(doc, "Evidence Summary", level=2, color=HEADER_BLUE)

    spans = evidence_result.get("evidence_spans", [])
    coverage = evidence_result.get("coverage_stats", {})

    stats_text = (
        f"Total evidence spans: {len(spans)}  |  "
        f"Nexus factors: {evidence_result.get('nexus_count', 0)}  |  "
        f"Note coverage: {coverage.get('coverage_pct', 0)}%"
    )
    _add_para(doc, stats_text, size=10)

    ungrounded = evidence_result.get("ungrounded_branches", [])
    if ungrounded:
        p = doc.add_paragraph()
        run = p.add_run(f"⚠ Ungrounded branches (no evidence found): {', '.join(ungrounded)}")
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        run.font.size = Pt(10)

    # Factor → Evidence → Branch table
    if spans:
        factor_table = doc.add_table(rows=1 + len(spans[:20]), cols=4)
        factor_table.style = 'Table Grid'

        headers = ["#", "Factor", "Evidence (from note)", "Branch(es)"]
        for i, h in enumerate(headers):
            cell = factor_table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

        for idx, span in enumerate(spans[:20], 1):
            row = factor_table.rows[idx]
            row.cells[0].text = str(idx)

            factors_str = "; ".join(span.get("factors", [])[:3])
            row.cells[1].text = factors_str[:60]

            evidence_text = span.get("text", "")[:100]
            row.cells[2].text = f'"{evidence_text}"' if evidence_text else "(no text)"

            branches = span.get("branches", [])
            branches_str = ", ".join(branches)
            row.cells[3].text = branches_str

            # Color the row based on branch
            color_key = span.get("color_key", "A")
            color_hex = BRANCH_COLORS.get(color_key, {}).get("hex", "FFFFFF")
            _set_cell_shading(row.cells[0], color_hex)


# ============================================================================
# PAGE 2-3: HIGHLIGHTED CLINICAL NOTE
# ============================================================================


def _build_highlighted_note_page(doc, note_text: str, evidence_result: dict):
    """
    Build the clinical note with color-highlighted evidence spans.
    """
    doc.add_page_break()
    _add_section_header(doc, "Clinical Note with Evidence Highlights", level=1, color=HEADER_BLUE)

    _add_para(
        doc,
        "Highlighted spans show which parts of the original clinical note "
        "ground each counterfactual scenario branch.",
        size=9,
    )
    doc.add_paragraph()

    # Build a sorted list of (start, end, color_hex) from evidence spans
    spans = evidence_result.get("evidence_spans", [])
    highlight_regions = []
    for s in spans:
        start = s.get("start", 0)
        end = s.get("end", 0)
        color_key = s.get("color_key", "A")
        color_hex = BRANCH_COLORS.get(color_key, {}).get("hex", "FFFF00")
        branches = s.get("branches", [])
        if start < end <= len(note_text):
            highlight_regions.append((start, end, color_hex, branches))

    # Sort by start position, remove overlaps (keep first)
    highlight_regions.sort(key=lambda x: x[0])
    cleaned_regions = []
    last_end = 0
    for start, end, color_hex, branches in highlight_regions:
        if start >= last_end:
            cleaned_regions.append((start, end, color_hex, branches))
            last_end = end
        elif start < last_end < end:
            # Partial overlap — take the non-overlapping tail
            cleaned_regions.append((last_end, end, color_hex, branches))
            last_end = end

    # Render note text with highlights
    # Split into paragraphs by double newlines
    # For each paragraph, find highlights within it and render runs
    para_splits = note_text.split('\n')

    current_pos = 0
    p = doc.add_paragraph()
    p.style.font.size = Pt(9)

    for line_idx, line in enumerate(para_splits):
        line_start = note_text.find(line, current_pos)
        if line_start == -1:
            line_start = current_pos
        line_end = line_start + len(line)

        # Find highlights in this line
        line_highlights = [
            (max(s, line_start) - line_start, min(e, line_end) - line_start, c, b)
            for s, e, c, b in cleaned_regions
            if s < line_end and e > line_start
        ]

        if not line.strip() and line_idx > 0:
            # Empty line — new paragraph
            p = doc.add_paragraph()
            p.style.font.size = Pt(9)
            current_pos = line_end + 1
            continue

        if not line_highlights:
            # No highlights — plain text
            run = p.add_run(line + "\n")
            run.font.size = Pt(9)
        else:
            # Render with highlights
            pos = 0
            for hl_start, hl_end, color_hex, branches in line_highlights:
                # Text before highlight
                if pos < hl_start:
                    run = p.add_run(line[pos:hl_start])
                    run.font.size = Pt(9)

                # Highlighted text
                hl_text = line[hl_start:hl_end]
                run = p.add_run(hl_text)
                run.font.size = Pt(9)
                run.bold = True
                _add_highlight(run, color_hex)

                # Add branch annotation as superscript
                branch_str = ",".join(branches)
                ann_run = p.add_run(f"[{branch_str}]")
                ann_run.font.size = Pt(6)
                ann_run.font.superscript = True
                ann_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

                pos = hl_end

            # Remaining text after last highlight
            if pos < len(line):
                run = p.add_run(line[pos:] + "\n")
                run.font.size = Pt(9)
            else:
                p.add_run("\n").font.size = Pt(9)

        current_pos = line_end + 1


# ============================================================================
# PAGE 4+: FULL PREPAREDNESS REPORT
# ============================================================================


def _build_report_pages(
    doc, report: dict, silver_labels: dict,
    scenarios_result: dict, critical_signals: dict,
):
    """Build the full preparedness report section."""
    doc.add_page_break()
    _add_section_header(doc, "Clinical Preparedness Report", level=1, color=HEADER_BLUE)

    # --- Critical Alerts ---
    all_alerts = []
    if critical_signals and critical_signals.get("signals_found"):
        all_alerts.extend(critical_signals.get("priority_alerts", []))
    all_alerts.extend(report.get("priority_alerts", []))
    # Deduplicate
    seen = set()
    unique_alerts = []
    for a in all_alerts:
        if a not in seen:
            seen.add(a)
            unique_alerts.append(a)

    if unique_alerts:
        _add_section_header(doc, "⚠ Critical Safety Alerts", level=2)
        for alert in unique_alerts:
            p = doc.add_paragraph()
            run = p.add_run(alert)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # --- Patient Overview ---
    _add_section_header(doc, "1. Patient Overview", level=2, color=HEADER_BLUE)
    _add_para(doc, report.get("patient_overview", "Not available."), size=10)

    # --- Mental Health Status ---
    _add_section_header(doc, "2. Mental Health Status", level=2, color=HEADER_BLUE)
    _add_para(doc, report.get("mental_health_status", "Not available."), size=10)

    # --- Key Risk Factors ---
    _add_section_header(doc, "3. Key Risk Factors", level=2, color=HEADER_BLUE)
    for rf in report.get("key_risk_factors", []):
        doc.add_paragraph(rf, style='List Bullet')

    # --- Protective Factors ---
    _add_section_header(doc, "4. Protective Factors", level=2, color=HEADER_BLUE)
    for pf in report.get("protective_factors", []):
        doc.add_paragraph(pf, style='List Bullet')

    # --- Scenario Briefing (branch-colored headers) ---
    _add_section_header(doc, "5. Counterfactual Scenario Briefing", level=2, color=HEADER_BLUE)

    for scenario in scenarios_result.get("scenarios", []):
        branch_id = scenario.get("scenario_id", "?")
        branch_label = scenario.get("branch_label", "?")
        color_info = BRANCH_COLORS.get(branch_id, BRANCH_COLORS["A"])

        # Branch header with color
        heading = doc.add_heading(f"[{branch_id}] {branch_label}", level=3)
        for run in heading.runs:
            run.font.color.rgb = color_info["rgb"]

        if scenario.get("gated"):
            _add_para(
                doc,
                f"NOT APPLICABLE: {scenario.get('gate_reason', 'Insufficient evidence')}",
                size=10,
            )
            continue

        scenario_obj = scenario.get("scenario", {})

        # Narrative
        narrative = scenario_obj.get("narrative", "")
        if narrative:
            _add_para(doc, narrative[:800], size=10)

        # Plausibility
        plaus = scenario_obj.get("plausibility", "?")
        p_rat = scenario_obj.get("plausibility_rationale", "")
        p = doc.add_paragraph()
        run = p.add_run(f"Plausibility: {plaus.upper()}")
        run.bold = True
        run.font.size = Pt(10)
        if p_rat:
            p.add_run(f" — {p_rat}").font.size = Pt(9)

        # Warning signs
        wsigns = scenario_obj.get("warning_signs", [])
        if wsigns:
            _add_para(doc, "Warning Signs:", bold=True, size=10)
            for ws in wsigns[:5]:
                doc.add_paragraph(ws, style='List Bullet')

        # Crisis endpoint
        crisis = scenario_obj.get("crisis_endpoint", "")
        if crisis:
            p = doc.add_paragraph()
            run = p.add_run("Crisis Endpoint: ")
            run.bold = True
            run.font.size = Pt(10)
            p.add_run(crisis).font.size = Pt(10)

        # Preparedness actions
        actions = scenario_obj.get("preparedness_actions", [])
        if actions:
            _add_para(doc, "Preparedness Actions:", bold=True, size=10)
            for a in actions[:5]:
                doc.add_paragraph(a, style='List Bullet')

    # --- Priority Actions ---
    _add_section_header(doc, "6. Priority Preparedness Actions", level=2, color=HEADER_BLUE)
    for i, action in enumerate(report.get("priority_actions", []), 1):
        doc.add_paragraph(f"{i}. {action}")

    # --- Risk Tier ---
    risk_tier = report.get("overall_risk_tier", "?")
    _add_section_header(doc, f"7. Overall Risk Tier: {risk_tier}", level=2, color=HEADER_BLUE)
    justification = report.get("risk_tier_justification", "")
    if justification:
        _add_para(doc, justification, size=10)


# ============================================================================
# MAIN FUNCTION
# ============================================================================


def generate_docx_report(
    note_text: str,
    evidence_result: dict,
    report: dict,
    silver_labels: dict,
    scenarios_result: dict,
    critical_signals: dict = None,
    output_path: Optional[str] = None,
    subject_id: str = "unknown",
    hadm_id: str = "unknown",
) -> Optional[str]:
    """
    Generate a DOCX report with color-highlighted clinical note and
    full preparedness report.

    Args:
        note_text: Original clinical note text.
        evidence_result: Phase 4 output (evidence attribution).
        report: Phase 3 output (preparedness report).
        silver_labels: Phase 1 output (extracted factors).
        scenarios_result: Phase 2 output (scenarios).
        critical_signals: Phase 1.6 output (critical signals).
        output_path: Where to save the .docx file. Auto-generated if None.
        subject_id: Patient subject ID for filename.
        hadm_id: Admission ID for filename.

    Returns:
        Path to the generated .docx file, or None if python-docx is missing.
    """
    if not HAS_DOCX:
        log.warning(
            "Phase 4.5: python-docx not installed. "
            "Install with: pip install python-docx"
        )
        return None

    log.info("Phase 4.5: Generating DOCX report with evidence highlights ...")

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Build document sections
    _build_legend_page(doc, evidence_result, silver_labels, report)
    _build_highlighted_note_page(doc, note_text, evidence_result)
    _build_report_pages(doc, report, silver_labels, scenarios_result, critical_signals)

    # Save
    if output_path is None:
        output_path = f"report_{subject_id}_{hadm_id}.docx"

    doc.save(output_path)
    log.info(f"Phase 4.5: DOCX report saved → {output_path}")
    return output_path
